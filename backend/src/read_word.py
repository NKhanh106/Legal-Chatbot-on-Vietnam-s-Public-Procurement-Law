"""
Module đọc Word documents (.doc, .docx) và chuyển đổi sang Markdown.
Tối ưu cho văn bản pháp luật tiếng Việt - KHÔNG dùng OCR (vì file Word gốc).
Sử dụng python-docx để giữ nguyên chính tả và cấu trúc.
"""
import os
import sys
import re
from pathlib import Path
from typing import Optional, Tuple
import logging

# Cấu hình logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)

# Get project root directory
PROJECT_ROOT = os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
DOCUMENTS_DIR = os.path.join(PROJECT_ROOT, "documents")
MARKDOWN_DIR = os.path.join(DOCUMENTS_DIR, "markdown")

# Tạo thư mục markdown nếu chưa có
os.makedirs(MARKDOWN_DIR, exist_ok=True)

def optimize_markdown_output(markdown_content: str) -> str:
    """
    Tối ưu markdown output sau khi đọc Word - LẤY TỪ read_pdf.py.
    Tối ưu cho văn bản pháp luật tiếng Việt.
    """
    if not markdown_content:
        return ""
    
    # Loại bỏ số trang và header/footer - cải thiện để bắt nhiều pattern hơn
    # Pattern 1: "Trang X" hoặc "Page X" với các ký tự phân cách
    markdown_content = re.sub(r"=+\s*Trang\s*\d+\s*=+", "", markdown_content, flags=re.IGNORECASE)
    markdown_content = re.sub(r"-\s*\d+\s*-", "", markdown_content)
    markdown_content = re.sub(r"^\s*Page\s+\d+\s*$", "", markdown_content, flags=re.MULTILINE | re.IGNORECASE)
    
    # Pattern 2: Số trang đứng một mình trên dòng (chỉ số, có thể có ký tự đặc biệt)
    markdown_content = re.sub(r"^\s*[^\w]*\d+[^\w]*\s*$", "", markdown_content, flags=re.MULTILINE)
    
    # Pattern 3: Số trang ở đầu dòng với ký tự đặc biệt
    markdown_content = re.sub(r"^[^\w]*\d+[^\w]*\s+", "", markdown_content, flags=re.MULTILINE)
    
    # Pattern 4: Số trang ở cuối dòng với ký tự đặc biệt
    markdown_content = re.sub(r"\s+[^\w]*\d+[^\w]*$", "", markdown_content, flags=re.MULTILINE)
    
    # Pattern 5: Dòng chỉ chứa số và ký tự đặc biệt (header/footer)
    markdown_content = re.sub(r"^[^\w\s]*\d+[^\w\s]*$", "", markdown_content, flags=re.MULTILINE)
    
    # Loại bỏ markdown artifacts
    markdown_content = re.sub(r"```[\s\S]*?```", "", markdown_content)
    markdown_content = re.sub(r"!\[.*?\]\(.*?\)", "", markdown_content)
    
    # Chuẩn hóa khoảng trắng và line endings
    markdown_content = re.sub(r"\r\n", "\n", markdown_content)
    markdown_content = re.sub(r"\r", "\n", markdown_content)
    markdown_content = re.sub(r"\n{3,}", "\n\n", markdown_content)
    markdown_content = re.sub(r"[ \t]+", " ", markdown_content)
    markdown_content = re.sub(r" +\n", "\n", markdown_content)
    
    # Chuẩn hóa headers markdown
    markdown_content = re.sub(r"^(#{1,6})([^\s#])", r"\1 \2", markdown_content, flags=re.MULTILINE)
    
    # Thêm phân cấp markdown cho cấu trúc pháp luật (quan trọng cho embedding)
    # Xử lý từng dòng để thêm headers phù hợp
    lines = markdown_content.split('\n')
    processed_lines = []
    
    for line in lines:
        line_stripped = line.strip()
        
        # Chương, Phần, Mục → ## (level 2)
        if re.match(r"^(Chương|Phần|Mục)\s+([IVX\d]+|[A-Z])\s*[\.:]?\s*", line_stripped, re.IGNORECASE):
            line_stripped = re.sub(
                r"^(Chương|Phần|Mục)\s+([IVX\d]+|[A-Z])\s*[\.:]?\s*(.+)?$",
                r"## \1 \2\3",
                line_stripped,
                flags=re.IGNORECASE
            )
            processed_lines.append(line_stripped)
            continue
        
        # Điều → ### (level 3) - chỉ khi chưa có header
        # Bao gồm cả chữ "đ" trong tiếng Việt (ví dụ: "Điều 1đ")
        if re.match(r"^Điều\s+\d+[a-zđ]?\s*[\.:]", line_stripped, re.IGNORECASE) and not line_stripped.startswith('#'):
            line_stripped = re.sub(
                r"^Điều\s+(\d+[a-zđ]?)\s*[\.:]\s*(.+)?$",
                r"### Điều \1. \2",
                line_stripped,
                flags=re.IGNORECASE
            )
            processed_lines.append(line_stripped)
            continue
        
        # Khoản (số) → #### (level 4) - chỉ khi là khoản của điều (số đơn giản)
        if re.match(r"^(\d+)\.\s+[A-ZĐ]", line_stripped) and len(line_stripped.split('.')[0].strip()) <= 2:
            line_stripped = re.sub(
                r"^(\d+)\.\s+(.+)$",
                r"#### \1. \2",
                line_stripped
            )
            processed_lines.append(line_stripped)
            continue
        
        # Điểm (a, b, c, d, đ, e, g...) → ##### (level 5)
        # Bao gồm cả chữ đ của tiếng Việt
        if re.match(r"^[a-zđ]\)\s+", line_stripped, re.IGNORECASE):
            line_stripped = re.sub(
                r"^([a-zđ])\)\s+(.+)$",
                r"##### \1) \2",
                line_stripped,
                flags=re.IGNORECASE
            )
            processed_lines.append(line_stripped)
            continue
        
        # Giữ nguyên các dòng khác
        processed_lines.append(line)
    
    markdown_content = '\n'.join(processed_lines)
    
    # Chuẩn hóa lại format sau khi thêm headers
    # Bao gồm cả chữ "đ" trong tiếng Việt
    markdown_content = re.sub(
        r"^(###\s+Điều\s+\d+[a-zđ]?)\s*\.\s*\.\s*",
        r"\1. ",
        markdown_content,
        flags=re.MULTILINE
    )
    
    # Loại bỏ ký tự control không hợp lệ
    markdown_content = re.sub(r"[\x00-\x08\x0B-\x0C\x0E-\x1F\x7F]", "", markdown_content)
    
    # Chuẩn hóa dấu ngoặc kép tiếng Việt
    markdown_content = markdown_content.replace('"', '"').replace('"', '"')
    markdown_content = markdown_content.replace("'", "'").replace("'", "'")
    
    # Loại bỏ dòng trống ở đầu và cuối
    lines = markdown_content.split('\n')
    while lines and not lines[0].strip():
        lines.pop(0)
    while lines and not lines[-1].strip():
        lines.pop()
    
    markdown_content = '\n'.join(lines)
    
    if markdown_content and not markdown_content.endswith('\n'):
        markdown_content += '\n'
    
    return markdown_content.strip()

def extract_text_with_formatting(paragraph) -> str:
    """
    Trích xuất text từ paragraph với formatting (bold, italic).
    Giữ nguyên chính tả tiếng Việt.
    """
    text_parts = []
    
    for run in paragraph.runs:
        run_text = run.text
        if not run_text:
            continue
        
        # Xử lý formatting
        if run.bold:
            text_parts.append(f"**{run_text}**")
        elif run.italic:
            text_parts.append(f"*{run_text}*")
        else:
            text_parts.append(run_text)
    
    return "".join(text_parts)

def iter_block_items(parent):
    """
    Lặp qua tất cả block (paragraph hoặc table) theo thứ tự trong Word document.
    Đảm bảo giữ nguyên thứ tự xuất hiện trong file Word.
    """
    from docx.oxml.text.paragraph import CT_P
    from docx.oxml.table import CT_Tbl
    
    parent_element = parent.element.body
    for element in parent_element.iterchildren():
        if isinstance(element, CT_P):
            # Tìm paragraph tương ứng
            for para in parent.paragraphs:
                if para._element == element:
                    yield ('paragraph', para)
                    break
        elif isinstance(element, CT_Tbl):
            # Tìm table tương ứng
            for table in parent.tables:
                if table._element == element:
                    yield ('table', table)
                    break

def is_list_item(paragraph) -> Tuple[bool, str]:
    """
    Kiểm tra xem paragraph có phải là list item không.
    Trả về (is_list, prefix) với prefix là "- " hoặc "1. " hoặc "".
    """
    style_name = paragraph.style.name.lower()
    
    # Kiểm tra style name
    if 'list' in style_name:
        if 'number' in style_name or 'numbered' in style_name:
            # Numbered list - cần đếm số thứ tự (tạm thời dùng "1. ")
            return (True, "1. ")
        else:
            # Bullet list
            return (True, "- ")
    
    # Kiểm tra paragraph format
    try:
        if paragraph._element.pPr is not None:
            numPr = paragraph._element.pPr.numPr
            if numPr is not None:
                # Có numbering properties - là list item
                ilvl = numPr.ilvl
                if ilvl is not None and ilvl.val is not None:
                    # Có level - là list item
                    # Kiểm tra loại list (bullet hay numbered)
                    numId = numPr.numId
                    if numId is not None:
                        # Có thể là numbered list
                        return (True, "1. ")
                    return (True, "- ")
    except (AttributeError, TypeError):
        pass
    
    # Kiểm tra text pattern (fallback)
    text = paragraph.text.strip()
    if re.match(r"^[\d]+[\.\)]\s+", text):
        return (True, "")  # Đã có số trong text
    if text.startswith("- ") or text.startswith("* ") or text.startswith("• "):
        return (True, "")  # Đã có bullet trong text
    
    return (False, "")

def read_docx_to_markdown(docx_path: str, output_path: str = None) -> str:
    """
    Đọc file .docx và chuyển đổi sang markdown.
    Sử dụng python-docx để giữ nguyên chính tả và cấu trúc.
    Cải tiến: Xử lý blocks theo thứ tự, cải thiện List/Header detection, xử lý bảng tốt hơn.
    
    Args:
        docx_path: Đường dẫn đến file .docx
        output_path: Đường dẫn file markdown output (nếu None thì tự động tạo)
    
    Returns:
        Đường dẫn đến file markdown đã tạo
    """
    try:
        from docx import Document
        from docx.oxml.text.paragraph import CT_P
        from docx.oxml.table import CT_Tbl
    except ImportError:
        logger.error("❌ Cần cài đặt python-docx: pip install python-docx")
        raise ImportError("python-docx chưa được cài đặt")
    
    docx_path = Path(docx_path)
    if not docx_path.exists():
        raise FileNotFoundError(f"❌ File không tồn tại: {docx_path}")
    
    if output_path is None:
        output_path = Path(MARKDOWN_DIR) / f"{docx_path.stem}.md"
    else:
        output_path = Path(output_path)
    
    logger.info(f"📄 Đang đọc Word: {docx_path.name}")
    logger.info(f"   Kích thước: {docx_path.stat().st_size / (1024*1024):.2f} MB")
    
    if output_path.exists():
        logger.info(f"   ⚠️  File markdown đã tồn tại, sẽ ghi đè: {output_path}")
    
    try:
        # Đọc document bằng python-docx (giữ nguyên chính tả)
        doc = Document(str(docx_path))
        
        # Chuyển đổi sang markdown - xử lý blocks theo thứ tự
        markdown_lines = []
        prev_was_heading = False
        prev_was_list = False
        prev_was_table = False
        
        # Xử lý tất cả blocks theo thứ tự (paragraphs và tables xen kẽ)
        for block_type, block in iter_block_items(doc):
            if block_type == 'paragraph':
                paragraph = block
                para_text = extract_text_with_formatting(paragraph)
                para_text = para_text.strip()
                
                if not para_text:
                    # Dòng trống - chỉ thêm nếu cần
                    if not prev_was_heading and not prev_was_table:
                        markdown_lines.append("")
                    prev_was_heading = False
                    prev_was_list = False
                    continue
                
                style_name = paragraph.style.name.lower()
                is_heading = False
                heading_level = 0
                
                # 1. Xử lý Headings - cải thiện detection
                if 'heading' in style_name:
                    # Phát hiện level từ style name
                    if 'heading 1' in style_name or 'heading1' in style_name:
                        heading_level = 1
                        is_heading = True
                    elif 'heading 2' in style_name or 'heading2' in style_name:
                        heading_level = 2
                        is_heading = True
                    elif 'heading 3' in style_name or 'heading3' in style_name:
                        heading_level = 3
                        is_heading = True
                    elif 'heading 4' in style_name or 'heading4' in style_name:
                        heading_level = 4
                        is_heading = True
                    elif 'heading 5' in style_name or 'heading5' in style_name:
                        heading_level = 5
                        is_heading = True
                    elif 'heading 6' in style_name or 'heading6' in style_name:
                        heading_level = 6
                        is_heading = True
                    else:
                        # Heading nhưng không rõ level - thử đoán từ style
                        if 'title' in style_name:
                            heading_level = 1
                            is_heading = True
                        elif 'subtitle' in style_name:
                            heading_level = 2
                            is_heading = True
                
                if is_heading and heading_level > 0:
                    if prev_was_heading or prev_was_table:
                        markdown_lines.append("")  # Dòng trống trước heading
                    markdown_lines.append(f"{'#' * heading_level} {para_text}")
                    prev_was_heading = True
                    prev_was_list = False
                    continue
                
                # 2. Xử lý Lists - cải thiện detection
                is_list, list_prefix = is_list_item(paragraph)
                if is_list:
                    if prev_was_heading:
                        markdown_lines.append("")  # Dòng trống sau heading
                    if list_prefix:
                        markdown_lines.append(f"{list_prefix}{para_text}")
                    else:
                        markdown_lines.append(para_text)  # Đã có prefix trong text
                    prev_was_list = True
                    prev_was_heading = False
                    continue
                
                # 3. Paragraph thường
                if prev_was_heading:
                    markdown_lines.append("")  # Dòng trống sau heading
                if prev_was_list and not prev_was_heading:
                    markdown_lines.append("")  # Dòng trống sau list
                
                markdown_lines.append(para_text)
                prev_was_heading = False
                prev_was_list = False
                
            elif block_type == 'table':
                table = block
                if prev_was_heading or prev_was_list:
                    markdown_lines.append("")  # Dòng trống trước table
                
                if table.rows:
                    # Lấy số cột tối đa từ hàng đầu tiên (cần cho merged cells)
                    num_cols = len(table.rows[0].cells)
                    
                    # Header Row
                    header_cells = []
                    for cell in table.rows[0].cells:
                        # Lấy text từ tất cả paragraphs trong cell
                        cell_text_parts = []
                        for para in cell.paragraphs:
                            para_text = " ".join([run.text for run in para.runs if run.text.strip()])
                            if para_text.strip():
                                cell_text_parts.append(para_text.strip())
                        cell_text = " ".join(cell_text_parts)
                        cell_text = cell_text.strip().replace('\n', ' ').replace('|', '\\|')
                        header_cells.append(cell_text)
                    
                    if header_cells:
                        markdown_lines.append("| " + " | ".join(header_cells) + " |")
                        markdown_lines.append("| " + " | ".join(["---"] * num_cols) + " |")
                        
                        # Data Rows
                        for row in table.rows[1:]:
                            cells = []
                            for cell in row.cells:
                                # Lấy text từ tất cả paragraphs trong cell
                                cell_text_parts = []
                                for para in cell.paragraphs:
                                    para_text = " ".join([run.text for run in para.runs if run.text.strip()])
                                    if para_text.strip():
                                        cell_text_parts.append(para_text.strip())
                                cell_text = " ".join(cell_text_parts)
                                cell_text = cell_text.strip().replace('\n', ' ').replace('|', '\\|')
                                cells.append(cell_text)
                            
                            # Xử lý merged cells: nếu số cells ít hơn num_cols, thêm empty cells
                            # Lưu ý: python-docx có hạn chế với merged cells, nhưng ta cố gắng xử lý
                            while len(cells) < num_cols:
                                cells.append("")
                            cells = cells[:num_cols]  # Đảm bảo không vượt quá
                            markdown_lines.append("| " + " | ".join(cells) + " |")
                
                markdown_lines.append("")  # Dòng trống sau table
                prev_was_table = True
                prev_was_heading = False
                prev_was_list = False
        
        # Ghép tất cả lại
        markdown_content = "\n".join(markdown_lines)
        
        # Tối ưu markdown với post-processing cho văn bản pháp luật
        logger.info("   🧹 Đang tối ưu markdown (post-processing cho văn bản pháp luật)...")
        original_size = len(markdown_content)
        markdown_content = optimize_markdown_output(markdown_content)
        optimized_size = len(markdown_content)
        
        # Lưu file
        output_path.parent.mkdir(parents=True, exist_ok=True)
        with open(output_path, "w", encoding="utf-8", newline='\n', buffering=8192) as f:
            f.write(markdown_content)
        
        # Thống kê
        char_count = len(markdown_content)
        line_count = markdown_content.count('\n')
        word_count = len(re.findall(r'\b\w+\b', markdown_content))
        size_reduction = original_size - optimized_size
        
        logger.info(f"   ✅ Hoàn tất!")
        logger.info(f"   📊 Số ký tự: {char_count:,} (giảm {size_reduction:,} ký tự)")
        logger.info(f"   📊 Số dòng: {line_count:,}")
        logger.info(f"   📊 Số từ: {word_count:,}")
        logger.info(f"   💾 Đã lưu: {output_path}")
        
        return str(output_path)
        
    except KeyboardInterrupt:
        logger.warning(f"   ⚠️  Đã dừng xử lý {docx_path.name} bởi người dùng")
        raise
    except Exception as e:
        logger.error(f"   ❌ Lỗi khi xử lý {docx_path.name}: {e}")
        import traceback
        logger.debug(traceback.format_exc())
        raise

def read_doc_to_markdown(doc_path: str, output_path: str = None) -> str:
    """
    Đọc file .doc (cũ) và chuyển đổi sang markdown.
    Sử dụng python-docx2txt (KHÔNG dùng OCR).
    
    Args:
        doc_path: Đường dẫn đến file .doc
        output_path: Đường dẫn file markdown output
    
    Returns:
        Đường dẫn đến file markdown đã tạo
    """
    doc_path = Path(doc_path)
    if not doc_path.exists():
        raise FileNotFoundError(f"❌ File không tồn tại: {doc_path}")
    
    if output_path is None:
        output_path = Path(MARKDOWN_DIR) / f"{doc_path.stem}.md"
    else:
        output_path = Path(output_path)
    
    logger.info(f"📄 Đang đọc Word (.doc): {doc_path.name}")
    logger.info(f"   Kích thước: {doc_path.stat().st_size / (1024*1024):.2f} MB")
    
    if output_path.exists():
        logger.info(f"   ⚠️  File markdown đã tồn tại, sẽ ghi đè: {output_path}")
    
    # Sử dụng python-docx2txt (KHÔNG dùng OCR)
    try:
        import docx2txt
        logger.info("   🔄 Đang đọc bằng python-docx2txt (giữ nguyên chính tả)...")
        text = docx2txt.process(str(doc_path))
        logger.info("   ✅ Đọc thành công")
    except ImportError:
        logger.error("❌ Cần cài đặt python-docx2txt: pip install python-docx2txt")
        logger.error("   KHÔNG dùng OCR vì file Word gốc - cần giữ nguyên chính tả")
        raise ImportError("python-docx2txt chưa được cài đặt")
    except Exception as e:
        logger.error(f"   ❌ Lỗi khi đọc file .doc: {e}")
        raise
    
    if not text or len(text.strip()) < 100:
        logger.warning("   ⚠️  File có vẻ rỗng hoặc không đọc được")
    
    # Chuyển đổi text thành markdown đơn giản
    lines = text.split('\n')
    markdown_lines = []
    
    for line in lines:
        line = line.strip()
        if not line:
            markdown_lines.append("")
            continue
        
        # Phát hiện heading đơn giản (dòng ngắn, in hoa, hoặc có số)
        if len(line) < 80 and (line.isupper() or re.match(r'^(Chương|Điều|Mục|Khoản)\s+\d+', line, re.IGNORECASE)):
            markdown_lines.append(f"## {line}")
        else:
            markdown_lines.append(line)
    
    markdown_content = "\n".join(markdown_lines)
    
    # Tối ưu markdown với post-processing cho văn bản pháp luật
    logger.info("   🧹 Đang tối ưu markdown (post-processing cho văn bản pháp luật)...")
    original_size = len(markdown_content)
    markdown_content = optimize_markdown_output(markdown_content)
    optimized_size = len(markdown_content)
    
    # Lưu file
    output_path.parent.mkdir(parents=True, exist_ok=True)
    with open(output_path, "w", encoding="utf-8", newline='\n', buffering=8192) as f:
        f.write(markdown_content)
    
    # Thống kê
    char_count = len(markdown_content)
    line_count = markdown_content.count('\n')
    word_count = len(re.findall(r'\b\w+\b', markdown_content))
    size_reduction = original_size - optimized_size
    
    logger.info(f"   ✅ Hoàn tất!")
    logger.info(f"   📊 Số ký tự: {char_count:,} (giảm {size_reduction:,} ký tự)")
    logger.info(f"   📊 Số dòng: {line_count:,}")
    logger.info(f"   📊 Số từ: {word_count:,}")
    logger.info(f"   💾 Đã lưu: {output_path}")
    
    return str(output_path)

def process_all_word_files(documents_dir: str = DOCUMENTS_DIR) -> list:
    """
    Xử lý tất cả file Word trong thư mục documents.
    
    Args:
        documents_dir: Thư mục chứa file Word
    
    Returns:
        List các file markdown đã tạo
    """
    documents_dir = Path(documents_dir)
    if not documents_dir.exists():
        raise FileNotFoundError(f"❌ Không tìm thấy thư mục: {documents_dir}")
    
    # Tìm tất cả file Word
    docx_files = list(documents_dir.glob("*.docx"))
    doc_files = list(documents_dir.glob("*.doc"))
    
    all_files = docx_files + doc_files
    
    if not all_files:
        logger.warning(f"⚠️  Không tìm thấy file Word nào trong: {documents_dir}")
        return []
    
    logger.info(f"🔍 Tìm thấy {len(all_files)} file Word")
    logger.info(f"📁 Sẽ lưu markdown vào: {MARKDOWN_DIR}")
    logger.info("=" * 60)
    
    processed_files = []
    
    for word_file in all_files:
        try:
            if word_file.suffix.lower() == '.docx':
                output_path = read_docx_to_markdown(str(word_file))
            elif word_file.suffix.lower() == '.doc':
                output_path = read_doc_to_markdown(str(word_file))
            else:
                continue
            
            processed_files.append(output_path)
        except Exception as e:
            logger.error(f"❌ Lỗi khi xử lý {word_file.name}: {e}")
            continue
    
    logger.info("=" * 60)
    logger.info(f"✅ Đã xử lý {len(processed_files)}/{len(all_files)} file")
    
    return processed_files

if __name__ == "__main__":
    import argparse
    
    parser = argparse.ArgumentParser(description="Đọc Word documents và chuyển đổi sang Markdown (KHÔNG dùng OCR)")
    parser.add_argument(
        "--documents-dir",
        type=str,
        default=DOCUMENTS_DIR,
        help=f"Thư mục chứa file Word (mặc định: {DOCUMENTS_DIR})"
    )
    parser.add_argument(
        "--file",
        type=str,
        default=None,
        help="Xử lý một file cụ thể"
    )
    
    args = parser.parse_args()
    
    try:
        if args.file:
            # Xử lý một file
            file_path = Path(args.file)
            if file_path.suffix.lower() == '.docx':
                read_docx_to_markdown(str(file_path))
            elif file_path.suffix.lower() == '.doc':
                read_doc_to_markdown(str(file_path))
            else:
                logger.error("❌ File phải là .doc hoặc .docx")
        else:
            # Xử lý tất cả file
            process_all_word_files(args.documents_dir)
    except KeyboardInterrupt:
        logger.info("\n⚠️  Đã dừng bởi người dùng")
        sys.exit(1)
    except Exception as e:
        logger.error(f"❌ Lỗi: {e}")
        import traceback
        traceback.print_exc()
        sys.exit(1)
